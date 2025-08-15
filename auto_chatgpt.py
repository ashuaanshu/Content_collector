import streamlit as st
import time
import pyperclip
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from docx import Document
from io import BytesIO

def run_auto_chatgpt_app():
    st.title("📨 ChatGPT Content Automation")

    # Step 1: User input
    query = st.text_input("💬 Type your instruction or question:")
    content = st.text_area("📋 Paste content to send along with the query:", height=250)

    if st.button("🚀 Send to ChatGPT and Get Output"):
        if not query or not content:
            st.warning("⚠️ Please enter both query and content.")
            return

        quer = f"{query}\n\n+ {content}"
        pyperclip.copy(quer)
        st.info("🌀 Launching Chrome...")

        options = webdriver.ChromeOptions()
        options.add_argument("--disable-blink-features=AutomationControlled")
        driver = webdriver.Chrome(options=options)

        try:
            driver.get("https://chat.qwen.ai/")
            time.sleep(3)  # Allow manual login if needed

            # Paste input
            textarea = driver.find_element(By.XPATH, "//textarea[@id='chat-input']")
            textarea.click()
            textarea.send_keys(Keys.CONTROL, 'v')
            time.sleep(1)
            textarea.send_keys(Keys.ENTER)
            st.info("⏳ Query sent. Waiting for full response...")

            # Wait for at least one response block
            WebDriverWait(driver, 60).until(
                EC.presence_of_element_located((By.CLASS_NAME, "svelte-1c06zsf"))
            )

            # Wait for response to stop growing
            last_text = ""
            same_count = 0
            max_wait = 60
            check_interval = 2
            start_time = time.time()

            while time.time() - start_time < max_wait:
                blocks = driver.find_elements(By.CLASS_NAME, "svelte-1c06zsf")
                if not blocks:
                    continue

                current_text = blocks[-1].text.strip()
                if current_text == last_text:
                    same_count += 1
                else:
                    same_count = 0
                    last_text = current_text

                if same_count >= 3:
                    break

                time.sleep(check_interval)

            # Get full response text
            output_blocks = driver.find_elements(By.CLASS_NAME, "svelte-1c06zsf")
            all_text = "\n\n".join([el.text for el in output_blocks if el.text.strip() != ""])

            # Display and save output
            if all_text:
                if st.button("💾 Save Output as .docx"):
                    doc = Document()
                    doc.add_heading("ChatGPT Response", level=1)
                    for paragraph in all_text.split("\n\n"):
                        doc.add_paragraph(paragraph)

                    buffer = BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)

                    st.download_button(
                        label="📥 Download .docx file",
                        data=buffer,
                        file_name="chatgpt_response.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                st.divider()
                st.subheader("📋 ChatGPT Response")
                st.text_area("Generated Output:", all_text, height=300)

        except Exception as e:
            st.error(f"❌ Error occurred: {e}")
        finally:
            driver.quit()
